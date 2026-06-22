from io import BytesIO


def test_parse_plain_text_document():
    from app.integrations.document_parser import parse_document_content

    assert parse_document_content("hello 文档".encode("utf-8"), "demo.txt") == "hello 文档"


def test_parse_docx_document():
    from docx import Document

    from app.integrations.document_parser import parse_document_content

    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    buffer = BytesIO()
    doc.save(buffer)

    parsed = parse_document_content(buffer.getvalue(), "demo.docx")

    assert "第一段" in parsed
    assert "第二段" in parsed


def test_parse_docx_heading_anchor():
    from docx import Document

    from app.integrations.document_parser import parse_document_content

    doc = Document()
    doc.add_heading("项目背景", level=1)
    doc.add_paragraph("正文内容")
    buffer = BytesIO()
    doc.save(buffer)

    parsed = parse_document_content(buffer.getvalue(), "demo.docx")

    assert "[[page:1;anchor:项目背景;heading:1]]" in parsed
    assert "[[page:1;anchor:项目背景]]" in parsed

