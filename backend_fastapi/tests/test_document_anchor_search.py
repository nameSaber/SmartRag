from io import BytesIO


def test_docx_upload_search_returns_heading_anchor(client):
    from docx import Document

    headers = _login(client, "anchor-user")
    doc = Document()
    doc.add_heading("项目背景", level=1)
    doc.add_paragraph("正文内容用于检索")
    buffer = BytesIO()
    doc.save(buffer)
    content = buffer.getvalue()

    client.post(
        "/api/v1/upload/chunk",
        headers=headers,
        data={"fileMd5": "anchor-docx", "chunkIndex": 0, "totalChunks": 1, "fileName": "anchor.docx", "totalSize": len(content)},
        files={"file": ("anchor.docx", content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    client.post("/api/v1/upload/merge", headers=headers, json={"fileMd5": "anchor-docx", "fileName": "anchor.docx", "totalChunks": 1})

    result = client.get("/api/v1/search/hybrid", headers=headers, params={"query": "正文内容", "topK": 3}).json()

    assert result["code"] == 200
    assert result["data"][0]["pageNumber"] == 1
    assert result["data"][0]["anchorText"] == "项目背景"


def _login(client, username):
    client.post("/api/v1/users/register", json={"username": username, "password": "secret"})
    token = client.post("/api/v1/users/login", json={"username": username, "password": "secret"}).json()["data"]["token"]
    return {"Authorization": f"Bearer {token}"}

